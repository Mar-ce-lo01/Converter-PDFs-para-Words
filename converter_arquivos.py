"""
PDF to Word Converter with GUI
=================================
A simple GUI application built with Tkinter that converts every PDF file
in a selected folder to Microsoft Word (.docx) format using PyMuPDF and
python‑docx.

Main Features
-------------
* Batch conversion (all PDFs found in the chosen source folder)
* Optional extraction of images and insertion into the generated DOCX
* Attempts to preserve tables discovered by PyMuPDF
* Runs the conversion in a background thread so the interface remains
  responsive
* Progress bar and status updates

Usage
-----
1. Install the dependencies (Python 3.9+ recommended)::

       pip install pymupdf python-docx

2. Run the script::

       python pdf_to_word_converter.py

3. In the GUI choose a *PDF folder* and a *DOCX destination folder* then
   click **Convert PDFs → DOCX**.

Author: ChatGPT
Licence: MIT
"""
from __future__ import annotations

import glob
import os
import threading
import tkinter as tk
from datetime import datetime
from pathlib import Path
from tkinter import filedialog, messagebox, ttk

import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches

# ---------------------------------------------------------------------------
# Helper functions
# ---------------------------------------------------------------------------

def extract_tables(page: fitz.Page, doc: Document) -> None:
    """Find tables on *page* (PyMuPDF >= 1.24) and copy them to *doc*."""
    try:
        tables = page.find_tables()
    except AttributeError:
        # Older PyMuPDF versions do not have find_tables
        tables = None
    if not tables or not tables.tables:
        return

    for table in tables.tables:
        data = table.extract()
        if not data:
            continue
        rows, cols = len(data), len(data[0])
        doc_table = doc.add_table(rows=rows, cols=cols)
        for r, row in enumerate(data):
            for c, cell in enumerate(row):
                doc_table.cell(r, c).text = (cell or "").strip()


def extract_images(page: fitz.Page, doc: Document, temp_dir: Path) -> None:
    """Extract all raster images from *page* and insert them into *doc*.

    The images are written temporarily to *temp_dir* and deleted by the
    caller afterwards.
    """
    for img in page.get_images(full=True):
        xref = img[0]
        base_name = f"img_{xref}.png"
        img_bytes = page.parent.extract_image(xref)["image"]
        tmp_path = temp_dir / base_name
        tmp_path.write_bytes(img_bytes)
        try:
            # You may adjust the width here as needed
            doc.add_picture(str(tmp_path), width=Inches(2.5))
        except Exception:  # corrupted image etc.
            pass


# ---------------------------------------------------------------------------
# GUI application
# ---------------------------------------------------------------------------

class PDFtoWordGUI(tk.Tk):
    """Main window of the application."""

    def __init__(self) -> None:
        super().__init__()
        self.title("Conversor PDF → Word")
        self.minsize(550, 300)

        # Tkinter variables
        self.src_dir = tk.StringVar(value=str(Path.cwd() / "pdfs"))
        self.dst_dir = tk.StringVar(value=str(Path.cwd() / "docs_word"))
        self.include_imgs = tk.BooleanVar(value=True)
        self.status = tk.StringVar(value="Pronto")

        # Build directory skeleton if absent
        Path(self.src_dir.get()).mkdir(parents=True, exist_ok=True)
        Path(self.dst_dir.get()).mkdir(parents=True, exist_ok=True)

        self._build_widgets()

    # ------------------------------------------------------------------
    #  GUI layout helpers
    # ------------------------------------------------------------------

    def _build_widgets(self) -> None:
        padding = {"padx": 6, "pady": 4}
        frm = ttk.Frame(self, padding=20)
        frm.pack(fill=tk.BOTH, expand=True)

        # --- Source folder ---------------------------------------------------
        ttk.Label(frm, text="Pasta com PDFs:").grid(row=0, column=0, sticky="w", **padding)
        ttk.Entry(frm, textvariable=self.src_dir, width=45).grid(row=0, column=1, **padding)
        ttk.Button(frm, text="Procurar…", command=self._choose_src).grid(row=0, column=2, **padding)

        # --- Destination folder ---------------------------------------------
        ttk.Label(frm, text="Salvar DOCX em:").grid(row=1, column=0, sticky="w", **padding)
        ttk.Entry(frm, textvariable=self.dst_dir, width=45).grid(row=1, column=1, **padding)
        ttk.Button(frm, text="Procurar…", command=self._choose_dst).grid(row=1, column=2, **padding)

        # --- Options ---------------------------------------------------------
        ttk.Checkbutton(frm, text="Incluir imagens", variable=self.include_imgs).grid(
            row=2, column=0, columnspan=3, sticky="w", **padding
        )

        # --- Convert button --------------------------------------------------
        ttk.Button(frm, text="Converter PDFs → DOCX", command=self._start_conversion).grid(
            row=3, column=0, columnspan=3, pady=(12, 4)
        )

        # --- Progressbar + status -------------------------------------------
        self.prog = ttk.Progressbar(frm, orient="horizontal", mode="determinate", length=380)
        self.prog.grid(row=4, column=0, columnspan=3, pady=(8, 4))
        ttk.Label(frm, textvariable=self.status).grid(row=5, column=0, columnspan=3, sticky="w", **padding)

    # ------------------------------------------------------------------
    #  Callbacks
    # ------------------------------------------------------------------

    def _choose_src(self) -> None:
        folder = filedialog.askdirectory(title="Selecione a pasta com PDFs")
        if folder:
            self.src_dir.set(folder)

    def _choose_dst(self) -> None:
        folder = filedialog.askdirectory(title="Selecione a pasta para DOCX")
        if folder:
            self.dst_dir.set(folder)

    def _start_conversion(self) -> None:
        """Spawn a worker thread so the UI stays responsive."""
        threading.Thread(target=self._convert_all_pdfs, daemon=True).start()

    # ------------------------------------------------------------------
    #  Conversion logic
    # ------------------------------------------------------------------

    def _convert_all_pdfs(self) -> None:
        pdf_paths = sorted(Path(self.src_dir.get()).glob("*.pdf"))
        total = len(pdf_paths)
        if not total:
            messagebox.showwarning("Aviso", "Nenhum PDF encontrado na pasta selecionada.")
            return

        # Prep progress bar
        self.prog.configure(maximum=total, value=0)
        self._set_status("Convertendo 0/%d…" % total)

        for idx, pdf_path in enumerate(pdf_paths, start=1):
            self._set_status(f"{pdf_path.name} ({idx}/{total})")
            self._convert_single(pdf_path)
            self._update_progress(idx)

        self._set_status(f"Concluído! {total} arquivo(s) convertidos.")
        messagebox.showinfo("Sucesso", "Conversão finalizada com êxito!")
        self.prog.after(2000, lambda: self.prog.configure(value=0))

    def _convert_single(self, pdf_path: Path) -> None:
        dst_file = Path(self.dst_dir.get()) / (pdf_path.stem + ".docx")
        try:
            convert_pdf_to_docx(
                pdf_path=pdf_path,
                docx_path=dst_file,
                include_images=self.include_imgs.get(),
            )
        except Exception as exc:
            messagebox.showerror("Erro", f"Falha ao converter {pdf_path.name}: {exc}")

    # ------------------------------------------------------------------
    #  Thread‑safe Tk helper wrappers
    # ------------------------------------------------------------------

    def _set_status(self, txt: str) -> None:
        self.status.set(txt)
        self.update_idletasks()

    def _update_progress(self, val: int) -> None:
        self.prog.configure(value=val)
        self.update_idletasks()


# ---------------------------------------------------------------------------
#  Core conversion engine (can be tested standalone)
# ---------------------------------------------------------------------------

def convert_pdf_to_docx(pdf_path: Path, docx_path: Path, *, include_images: bool = True) -> None:
    """Convert *pdf_path* → *docx_path* preserving text, tables and images."""

    pdf = fitz.open(pdf_path)
    doc = Document()

    temp_dir = docx_path.with_name("_tmp_imgs_%s" % datetime.now().strftime("%Y%m%d%H%M%S"))
    temp_dir.mkdir(exist_ok=True)

    try:
        for page in pdf:
            text = page.get_text("text").strip()
            if text:
                doc.add_paragraph(text)

            extract_tables(page, doc)

            if include_images:
                extract_images(page, doc, temp_dir)

        doc.save(docx_path)
    finally:
        # Clean‑up temporary image files
        for img_file in temp_dir.glob("*"):
            try:
                img_file.unlink(missing_ok=True)
            except Exception:
                pass
        try:
            temp_dir.rmdir()
        except OSError:
            pass
        pdf.close()


# ---------------------------------------------------------------------------
#  Main entry point
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    app = PDFtoWordGUI()
    app.mainloop()
